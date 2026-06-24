"""
DOCX Builder Service — exports documents to Word format using python-docx.

Maps Document + DocumentSections to a styled Word document with
brand fonts, colors, and embedded images.
"""
import io
import logging

logger = logging.getLogger(__name__)


def build_docx(document):
    """
    Build a Word document from a Document model instance.

    Returns an io.BytesIO buffer containing the .docx file.
    """
    try:
        from docx import Document as DocxDocument
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError(
            'python-docx is not installed. Run: pip install python-docx>=1.1.0'
        )

    docx = DocxDocument()
    sections = document.sections.filter(is_visible=True).order_by('order')

    # Apply brand styling
    brand = document.brand_guideline
    _apply_brand_styles(docx, brand)

    # Title page
    title_para = docx.add_heading(document.title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if document.recipient_name:
        recipient_para = docx.add_paragraph(f'Prepared for: {document.recipient_name}')
        recipient_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    docx.add_page_break()

    # Sections
    for section in sections:
        if section.title:
            heading_level = 1 if section.section_type in ('COVER', 'EXECUTIVE_SUMMARY') else 2
            docx.add_heading(section.title, level=heading_level)

        if section.content_html:
            try:
                from htmldocx import HtmlToDocx
                new_parser = HtmlToDocx()
                # Wrap the HTML to ensure it's a valid fragment, just in case
                wrapped_html = f"<div>{section.content_html}</div>"
                new_parser.add_html_to_document(wrapped_html, docx)
            except Exception as e:
                logger.error(f'Could not convert HTML to DOCX using htmldocx: {e}')
                # Fallback to plain text
                clean_text = _strip_html(section.content_html)
                if clean_text:
                    docx.add_paragraph(clean_text)

        # Add image if present
        if section.image:
            try:
                docx.add_picture(section.image.path, width=Inches(5.0))
            except Exception as e:
                logger.warning(f'Could not embed image in DOCX: {e}')

    # Save to buffer
    buffer = io.BytesIO()
    docx.save(buffer)
    buffer.seek(0)
    return buffer


def _apply_brand_styles(docx, brand):
    """Apply brand colors and fonts to the DOCX styles."""
    if not brand:
        return

    try:
        from docx.shared import RGBColor, Pt

        style = docx.styles['Title']
        if brand.heading_font:
            style.font.name = brand.heading_font
        if brand.primary_color:
            style.font.color.rgb = _hex_to_rgb(brand.primary_color)

        for heading_level in ['Heading 1', 'Heading 2', 'Heading 3']:
            if heading_level in docx.styles:
                h_style = docx.styles[heading_level]
                if brand.heading_font:
                    h_style.font.name = brand.heading_font
                if brand.primary_color:
                    h_style.font.color.rgb = _hex_to_rgb(brand.primary_color)

        # Body text
        normal_style = docx.styles['Normal']
        if brand.body_font:
            normal_style.font.name = brand.body_font
        if brand.text_color:
            normal_style.font.color.rgb = _hex_to_rgb(brand.text_color)
        normal_style.font.size = Pt(11)

    except Exception as e:
        logger.warning(f'Could not apply brand styles to DOCX: {e}')


def _hex_to_rgb(hex_color):
    """Convert hex color to python-docx RGBColor."""
    from docx.shared import RGBColor
    hex_color = hex_color.lstrip('#')
    return RGBColor(
        int(hex_color[0:2], 16),
        int(hex_color[2:4], 16),
        int(hex_color[4:6], 16),
    )


def _strip_html(html_content):
    """Simple HTML tag stripping for DOCX text content."""
    import re
    # Remove HTML tags
    clean = re.sub(r'<[^>]+>', '', html_content)
    # Decode common HTML entities
    clean = clean.replace('&amp;', '&').replace('&lt;', '<').replace('&gt;', '>')
    clean = clean.replace('&nbsp;', ' ').replace('&quot;', '"')
    return clean.strip()
